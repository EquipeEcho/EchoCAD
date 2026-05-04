"""Exemplos de Integração do Módulo Knowledge com EchoCAD

Este arquivo mostra como integrar o agente RAG de normas em diferentes partes
do EchoCAD para enriquecer especificações técnicas, memorials, etc.
"""

# ============================================================================
# EXEMPLO 1: Integração com Gerador de Especificações
# ============================================================================

"""
Em: src/modules/EspecificacoesTecnicas/spec_generator.py

Adicionar consulta de normas ao gerar especificações:
"""

from results.knowledge import query_normas


def gerar_especificacoes_com_normas(
    project_id: int,
    projeto_nome: str,
    disciplinas: list[str],
) -> dict:
    """Gera especificações técnicas enriquecidas com contexto de normas."""
    
    especificacoes = {}
    
    for disciplina in disciplinas:
        # Consultar normas relevantes
        norma_query = f"Requisitos técnicos para {disciplina} conforme normas ABNT"
        contexto_normas = query_normas(norma_query)
        
        # Usar contexto para gerar especificação
        espec = {
            "disciplina": disciplina,
            "contexto_normas": contexto_normas,
            "conteudo": f"Especificações baseadas em: {contexto_normas[:500]}..."
        }
        
        especificacoes[disciplina] = espec
    
    return especificacoes


# ============================================================================
# EXEMPLO 2: Endpoint FastAPI com Contexto de Normas
# ============================================================================

"""
Em: src/routes/especificacoes.py

Adicionar rota que combina especificações com normas:
"""

from fastapi import APIRouter
from results.knowledge import query_normas

router = APIRouter()


@router.get("/especificacoes/{project_id}/com-normas")
async def get_especificacoes_com_normas(project_id: int):
    """Retorna especificações enriquecidas com contexto de normas."""
    
    # Obter especificações do projeto
    especificacoes = get_project_specifications(project_id)
    
    # Para cada disciplina, adicionar contexto de normas
    for espec in especificacoes:
        pergunta = f"Normas técnicas para {espec['disciplina']}"
        normas = query_normas(pergunta)
        espec["normas_aplicaveis"] = normas
    
    return especificacoes


# ============================================================================
# EXEMPLO 3: Validação de Projetos com Normas
# ============================================================================

"""
Em: src/routes/projeto_router.py

Adicionar validação de projeto conforme normas:
"""

from results.knowledge import query_normas


def validar_projeto_conforme_normas(projeto_id: int) -> dict:
    """Valida projeto contra normas técnicas."""
    
    projeto = get_projeto(projeto_id)
    validacoes = []
    
    # Validar estrutura
    norma_estrutura = query_normas(f"NBR 6118 - Verificação de estrutura {projeto.tipo}")
    validacoes.append({
        "area": "estrutura",
        "norma": "NBR 6118",
        "resultado": "verificado",
        "detalhes": norma_estrutura,
    })
    
    # Validar segurança
    norma_seguranca = query_normas("NR 18 - Segurança em construção")
    validacoes.append({
        "area": "seguranca",
        "norma": "NR 18",
        "resultado": "verificado",
        "detalhes": norma_seguranca,
    })
    
    return {
        "project_id": projeto_id,
        "validacoes": validacoes,
        "status": "validado",
    }


# ============================================================================
# EXEMPLO 4: Memorial de Cálculo com Referências a Normas
# ============================================================================

"""
Em: src/modules/core/main.py

Adicionar referências de normas ao memorial:
"""

from results.knowledge import query_normas


def executar_extracao_com_normas(
    prompt_usuario: str,
    dxf_path: str,
    project_id: int,
) -> dict:
    """Executa extração com referências às normas aplicáveis."""
    
    # Extração padrão
    resultado = run_extraction(prompt_usuario, dxf_path, project_id)
    
    # Adicionar contexto de normas ao resultado
    for item in resultado.get("items", []):
        # Consultar norma relevante para cada item
        norma = query_normas(f"Norma para {item['tipo']} em {item['categoria']}")
        item["norma_aplicavel"] = norma[:200]  # Primeiros 200 chars
    
    resultado["normas_consultadas"] = True
    
    return resultado


# ============================================================================
# EXEMPLO 5: Sistema de Recomendações Baseado em Normas
# ============================================================================

"""
Em: src/routes/projeto_router.py

Recomendar melhorias baseado em normas:
"""

from results.knowledge import query_normas


@router.get("/projeto/{project_id}/recomendacoes-normas")
async def get_recomendacoes_normas(project_id: int):
    """Gera recomendações de melhorias conforme normas."""
    
    projeto = get_projeto(project_id)
    recomendacoes = []
    
    # Verificar contra principais normas
    normas_principais = [
        "NBR 6118 - Estrutura de concreto",
        "NBR 14931 - Execução de estruturas",
        "NR 18 - Segurança em construção",
        "NBR 15968 - Coordenação modular",
    ]
    
    for norma in normas_principais:
        pergunta = f"Recomendações e requisitos: {norma}"
        resposta = query_normas(pergunta)
        
        recomendacoes.append({
            "norma": norma,
            "recomendacoes": resposta,
            "prioridade": "alta" if "deve" in resposta.lower() else "média",
        })
    
    return {
        "project_id": project_id,
        "recomendacoes": recomendacoes,
    }


# ============================================================================
# EXEMPLO 6: Chatbot Técnico para Consultas de Normas
# ============================================================================

"""
Em: src/routes/ (novo arquivo: chat_normas_router.py)

Chatbot para responder perguntas sobre normas:
"""

from fastapi import WebSocket
from results.knowledge import query_normas


@router.websocket("/ws/chat-normas/{project_id}")
async def websocket_chat_normas(websocket: WebSocket, project_id: int):
    """WebSocket para chat sobre normas técnicas."""
    
    await websocket.accept()
    projeto = get_projeto(project_id)
    
    try:
        while True:
            # Receber pergunta
            data = await websocket.receive_json()
            pergunta = data.get("pergunta", "")
            
            if not pergunta:
                await websocket.send_json({"erro": "Pergunta vazia"})
                continue
            
            # Consultar normas
            resposta = query_normas(pergunta)
            
            # Enviar resposta
            await websocket.send_json({
                "pergunta": pergunta,
                "resposta": resposta,
                "projeto": projeto.name,
                "timestamp": datetime.now().isoformat(),
            })
    
    except Exception as e:
        await websocket.send_json({"erro": str(e)})
    finally:
        await websocket.close()


# ============================================================================
# EXEMPLO 7: Geração de Relatório com Normas
# ============================================================================

"""
Em: src/routes/download.py

Gerar relatório que cita normas:
"""

from results.knowledge import query_normas
from docx import Document


def gerar_relatorio_com_normas(project_id: int) -> str:
    """Gera relatório técnico com citações de normas."""
    
    projeto = get_projeto(project_id)
    doc = Document()
    
    # Cabeçalho
    doc.add_heading(f"Relatório Técnico - {projeto.name}", level=1)
    
    # Seções
    secoes = [
        "Estrutura e Fundações",
        "Segurança em Construção",
        "Acabamentos",
        "Instalações Prediais",
    ]
    
    for secao in secoes:
        doc.add_heading(secao, level=2)
        
        # Consultar norma
        norma_texto = query_normas(f"Requisitos para {secao}")
        doc.add_paragraph(norma_texto)
    
    # Salvar
    filename = f"/tmp/relatorio_{project_id}.docx"
    doc.save(filename)
    
    return filename


# ============================================================================
# EXEMPLO 8: Cache de Consultas de Normas
# ============================================================================

"""
Em: src/modules/knowledge/normas_cache.py (novo arquivo)

Implementar cache de consultas para otimizar:
"""

from functools import lru_cache
from results.knowledge import query_normas


class NormasCache:
    """Cache para consultas de normas."""
    
    def __init__(self, max_size: int = 100):
        self.cache = {}
        self.max_size = max_size
    
    def query(self, pergunta: str) -> str:
        """Consulta com cache."""
        
        # Verificar cache
        if pergunta in self.cache:
            return self.cache[pergunta]
        
        # Executar consulta
        resposta = query_normas(pergunta)
        
        # Armazenar em cache
        if len(self.cache) >= self.max_size:
            # Remove entrada mais antiga
            oldest_key = next(iter(self.cache))
            del self.cache[oldest_key]
        
        self.cache[pergunta] = resposta
        
        return resposta
    
    def clear(self):
        """Limpa cache."""
        self.cache.clear()


# Uso global
_normas_cache = NormasCache()


def query_normas_cached(pergunta: str) -> str:
    """Consulta normas com cache automático."""
    return _normas_cache.query(pergunta)


# ============================================================================
# EXEMPLO 9: Integração com AI Orchestrator
# ============================================================================

"""
Em: src/modules/ai_orchestrator.py

Adicionar normas ao orquestrador de IA:
"""

from results.knowledge import query_normas


class AIOrchestrator:
    """Orquestrador com suporte a normas."""
    
    def processar_com_normas(self, documento: str, tipo: str):
        """Processa documento considerando normas."""
        
        # Obter contexto de normas
        norma_contexto = query_normas(f"Normas aplicáveis a {tipo}")
        
        # Combinar com agentes
        prompt_com_contexto = f"""
Processe o seguinte documento considerando:
CONTEXTO DE NORMAS:
{norma_contexto}

DOCUMENTO:
{documento}
"""
        
        # Executar pipeline com contexto
        resultado = self.execute_pipeline(prompt_com_contexto)
        
        return resultado


# ============================================================================
# EXEMPLO 10: Notificações sobre Normas Atualizadas
# ============================================================================

"""
Em: src/routes/notificacoes_router.py

Notificar sobre mudanças em normas:
"""

from results.knowledge import query_normas
from datetime import datetime


async def notificar_normas_atualizadas(project_id: int):
    """Notifica quando normas aplicáveis são atualizadas."""
    
    projeto = get_projeto(project_id)
    
    # Normas do projeto
    normas_do_projeto = [
        "NBR 6118",
        "NBR 14931",
        "NR 18",
    ]
    
    # Verificar cada norma
    for norma_codigo in normas_do_projeto:
        info_norma = query_normas(f"Informações sobre atualização de {norma_codigo}")
        
        # Se encontrou info sobre atualização
        if "atualização" in info_norma.lower() or "versão" in info_norma.lower():
            # Enviar notificação
            await enviar_notificacao(
                usuario_id=projeto.id_user,
                titulo=f"Norma {norma_codigo} Atualizada",
                mensagem=info_norma,
                tipo="norma_atualizada",
            )


# ============================================================================
# PADRÃO DE USO RECOMENDADO
# ============================================================================

"""
Para qualquer integração, siga este padrão:

1. Identifique o contexto técnico
   → Qual norma é aplicável?
   
2. Consulte o agente
   → resposta = query_normas("Pergunta técnica específica")
   
3. Use o contexto
   → Enriqueça dados, valide, recomende
   
4. Armazene referências
   → Mantenha rastreabilidade de qual norma foi usada
   
5. Permita auditoria
   → Mostre ao usuário quais normas foram consultadas

Exemplo:
    # 1. Contexto
    disciplina = "Estrutura"
    
    # 2. Consulta
    norma = query_normas(f"NBR 6118 - requisitos para {disciplina}")
    
    # 3. Use
    validar_contra_norma(projeto, norma)
    
    # 4. Armazene
    db.save_norma_referencia(projeto, "NBR 6118")
    
    # 5. Auditoria
    retornar {
        "resultado": "validado",
        "norma_aplicada": "NBR 6118",
        "referencia": norma[:200]
    }
"""

# ============================================================================
# ENDPOINTS RECOMENDADOS PARA ADICIONAR
# ============================================================================

"""
POST   /api/normas/consultar              # Query simples
POST   /api/normas/consultar-com-fontes  # Query com metadata
POST   /api/normas/ingest-arquivo        # Carregar PDF
POST   /api/normas/ingest-lote           # Carregar múltiplos
GET    /api/normas/status                # Status do módulo
GET    /api/normas/health                # Healthcheck

POST   /api/especificacoes/{id}/com-normas         # 🆕 Novo
GET    /api/projeto/{id}/recomendacoes-normas    # 🆕 Novo
GET    /api/projeto/{id}/validacao-normas        # 🆕 Novo
WS     /ws/chat-normas/{id}                      # 🆕 Novo
GET    /api/relatorio/{id}/com-normas            # 🆕 Novo
"""
