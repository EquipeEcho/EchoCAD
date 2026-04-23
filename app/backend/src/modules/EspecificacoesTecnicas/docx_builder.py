# docx_builder.py
# Converte o objeto EspecificacoesTecnicas em um documento Word (.docx)
# formatado seguindo o padrão do caderno de encargos do Exército.

import io
import logging
from pathlib import Path
from typing import List, Optional

from .spec_generator import EspecificacoesTecnicas, SecaoEspec

logger = logging.getLogger(__name__)


def _build_docx(specs: EspecificacoesTecnicas, output_path: str) -> Path:
    """
    Constrói o documento Word usando python-docx.
    Fallback para docx-js via Node se python-docx não estiver disponível.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return _build_with_python_docx(specs, output_path)
    except ImportError:
        logger.warning("python-docx não encontrado. Tentando alternativa...")
        return _build_with_xml(specs, output_path)


def _set_heading_style(paragraph, level: int, doc):
    """Aplica estilo de título ao parágrafo."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)

    if level == 0:  # Título principal
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:  # Seção principal (1., 2., ...)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    elif level == 2:  # Subseção (1.1, 1.2, ...)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    else:  # Nível mais profundo
        run.font.size = Pt(10)
        run.font.bold = True


def _add_table(doc, headers: List[str], rows: List[List[str]]):
    """Adiciona uma tabela formatada ao documento."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else \
              hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        # Fundo azul escuro no cabeçalho
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tc._new_tcShd()
        shd.set(qn('w:fill'), '003366')
        shd.set(qn('w:color'), 'FFFFFF')
        tcPr.append(shd)
        # Texto branco
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            row_cells[j].text = str(cell_text)
            p = row_cells[j].paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(cell_text))
            run.font.size = Pt(9)
            # Linhas alternadas
            if i % 2 == 0:
                tc = row_cells[j]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tc._new_tcShd()
                shd.set(qn('w:fill'), 'F5F5F5')
                tcPr.append(shd)


def _add_conteudo_formatado(doc, conteudo: str):
    """Adiciona texto com marcadores **bold** ao documento."""
    from docx.shared import Pt
    import re

    for linha in conteudo.split('\n'):
        linha = linha.rstrip()
        if not linha:
            doc.add_paragraph()
            continue

        # Detectar marcadores de lista
        if linha.startswith('- ') or linha.startswith('• '):
            p = doc.add_paragraph(style='List Bullet')
            _add_run_formatado(p, linha[2:])
        elif re.match(r'^\d+\. ', linha):
            p = doc.add_paragraph(style='List Number')
            _add_run_formatado(p, re.sub(r'^\d+\. ', '', linha))
        else:
            p = doc.add_paragraph()
            _add_run_formatado(p, linha)

        for run in p.runs:
            run.font.size = Pt(10)


def _add_run_formatado(paragraph, texto: str):
    """Adiciona runs ao parágrafo respeitando **negrito**."""
    import re
    from docx.shared import Pt

    partes = re.split(r'\*\*(.+?)\*\*', texto)
    for i, parte in enumerate(partes):
        if not parte:
            continue
        run = paragraph.add_run(parte)
        run.font.bold = (i % 2 == 1)
        run.font.size = Pt(10)


def _build_with_python_docx(specs: EspecificacoesTecnicas, output_path: str) -> Path:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ---- Configurar página ----
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)

    # ---- Página de capa ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Anexo 2 ao Projeto Básico")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Número Único de Protocolo: {specs.numero_protocolo}")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CADERNO DE ENCARGO E ESPECIFICAÇÕES TÉCNICAS")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(specs.nome_projeto.upper())
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_page_break()

    # ---- Finalidade ----
    h = doc.add_heading("FINALIDADE", level=1)
    _set_heading_style(h, 1, doc)
    doc.add_paragraph(
        "As presentes especificações técnicas têm por finalidade descrever os serviços "
        "a serem executados pela contratada, de modo que ela possa fornecer a mão de obra "
        "especializada, os materiais especificados e os equipamentos necessários à execução do objeto."
    ).runs[0].font.size = Pt(10)

    # ---- Objeto ----
    h = doc.add_heading("OBJETO", level=1)
    _set_heading_style(h, 1, doc)
    doc.add_paragraph(specs.objeto).runs[0].font.size = Pt(10)

    # ---- Acrônimos ----
    h = doc.add_heading("ACRÔNIMOS E SÍMBOLOS", level=1)
    _set_heading_style(h, 1, doc)
    _add_table(doc,
        ["Sigla", "Significado"],
        [
            ["ABNT", "Associação Brasileira de Normas Técnicas"],
            ["ART",  "Anotação de Responsabilidade Técnica"],
            ["CAU",  "Conselho de Arquitetura e Urbanismo"],
            ["CREA", "Conselho Regional de Engenharia e Agronomia"],
            ["DRT",  "Delegacia Regional do Trabalho"],
            ["INMETRO", "Instituto Nacional de Metrologia, Qualidade e Tecnologia"],
            ["NBR",  "Normas da ABNT"],
            ["NR",   "Normas Regulamentadoras do Ministério do Trabalho"],
            ["PCMAT","Programa de Condições e Meio Ambiente de Trabalho"],
            ["RRT",  "Registro de Responsabilidade Técnica"],
            ["SPDA", "Sistema de Proteção Contra Descargas Atmosféricas"],
        ]
    )
    doc.add_paragraph()

    # ---- Referências Normativas Gerais ----
    h = doc.add_heading("REFERÊNCIAS NORMATIVAS", level=1)
    _set_heading_style(h, 1, doc)
    for ref in specs.referencias_normativas:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ---- Vida Útil ----
    if specs.vida_util:
        h = doc.add_heading("VIDA ÚTIL E GARANTIAS", level=1)
        _set_heading_style(h, 1, doc)
        rows_vu = [
            [item.get('item',''), item.get('vida_util_anos',''),
             item.get('garantia_anos',''), item.get('nbr','')]
            for item in specs.vida_util
        ]
        _add_table(doc, ["Item", "Vida Útil [anos]", "Garantia [anos]", "NBR"], rows_vu)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- Seções principais ----
    for secao in specs.secoes:
        # Título da seção
        h = doc.add_heading(f"{secao.numero}. {secao.titulo}", level=1)
        _set_heading_style(h, 1, doc)

        # Introdução
        if secao.conteudo:
            _add_conteudo_formatado(doc, secao.conteudo)
            doc.add_paragraph()

        # Subseções
        for sub in secao.subsecoes:
            h2 = doc.add_heading(f"{sub.numero} {sub.titulo}", level=2)
            _set_heading_style(h2, 2, doc)

            if sub.conteudo:
                _add_conteudo_formatado(doc, sub.conteudo)

            doc.add_paragraph()

        doc.add_paragraph()

    # ---- Salvar ----
    out_p = Path(output_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    logger.info(f"Documento gerado: {out_p}")
    return out_p


def _build_with_xml(specs: EspecificacoesTecnicas, output_path: str) -> Path:
    """Fallback: gera um arquivo .txt estruturado se python-docx não estiver disponível."""
    out_p = Path(output_path).with_suffix('.txt')
    out_p.parent.mkdir(parents=True, exist_ok=True)

    lines = []
    lines.append("CADERNO DE ENCARGO E ESPECIFICAÇÕES TÉCNICAS")
    lines.append(f"Projeto: {specs.nome_projeto}")
    lines.append(f"Protocolo: {specs.numero_protocolo}")
    lines.append("=" * 70)
    lines.append("")
    lines.append("OBJETO")
    lines.append(specs.objeto)
    lines.append("")

    for secao in specs.secoes:
        lines.append(f"\n{'='*70}")
        lines.append(f"{secao.numero}. {secao.titulo.upper()}")
        lines.append('='*70)
        if secao.conteudo:
            lines.append(secao.conteudo)
        for sub in secao.subsecoes:
            lines.append(f"\n  {sub.numero} {sub.titulo}")
            lines.append("  " + "-" * 60)
            if sub.conteudo:
                for l in sub.conteudo.split('\n'):
                    lines.append("  " + l)

    out_p.write_text('\n'.join(lines), encoding='utf-8')
    logger.info(f"Documento (txt fallback) gerado: {out_p}")
    return out_p


def build_docx(specs: EspecificacoesTecnicas, output_path: str) -> Path:
    """Ponto de entrada público para construção do documento."""
    return _build_docx(specs, output_path)